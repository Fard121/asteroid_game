"""
Generates the AsteroidsFX technical report as a .docx following the structure
required by "Report template.pdf" and the deliverables listed in "info.docx".

Formatting: Times New Roman 12 pt, 1.5 line spacing, justified body text,
numbered headings, captioned figures, IEEE-style references.

Every figure and screenshot referenced here was generated from the actual
repository: the diagrams from Mermaid sources in docs/report-figures/*.mmd,
the screenshots from real runs of the built application.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
OUT = os.path.abspath(os.path.join(FIG, "..", "AsteroidsFX-Technical-Report.docx"))

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)
GREY = RGBColor(0x55, 0x55, 0x55)

fig_no = {"n": 0}
tab_no = {"n": 0}


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------
def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_field(paragraph, instr):
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = paragraph.add_run("1")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r4._r.append(end)


def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, colour, before, after in [
        ("Heading 1", 16, ACCENT, 18, 8),
        ("Heading 2", 13.5, ACCENT, 12, 6),
        ("Heading 3", 12.5, ACCENT, 10, 4),
    ]:
        s = doc.styles[name]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = colour
        s.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)


def P(doc, text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      space_after=6, colour=None, spacing=WD_LINE_SPACING.ONE_POINT_FIVE):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = spacing
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        if colour is not None:
            r.font.color.rgb = colour
    return p


def rich(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    """parts = list of (text, style) where style in {'', 'b', 'i', 'c'}"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for text, style in parts:
        r = p.add_run(text)
        if style == "b":
            r.bold = True
        elif style == "i":
            r.italic = True
        elif style == "c":
            r.font.name = MONO_FONT
            r.font.size = Pt(10.5)
            continue
        r.font.name = BODY_FONT
        r.font.size = Pt(12)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + 0.7 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(12)
    return p


def H(doc, text, level=1):
    return doc.add_heading(text, level=level)


def code(doc, text, caption=None):
    lines = text.strip("\n").split("\n")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F5F7")
    cell.paragraphs[0].text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
    if caption:
        cap(doc, caption, kind="Listing")


def cap(doc, text, kind="Figure"):
    if kind == "Figure":
        fig_no["n"] += 1
        n = fig_no["n"]
    else:
        tab_no["n"] += 1
        n = tab_no["n"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(f"{kind} {n}. {text}")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GREY
    return n


def figure(doc, filename, caption, width_cm=15.5):
    path = os.path.join(FIG, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    return cap(doc, caption)


def table(doc, headers, rows, caption, widths=None, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "E8ECF5")
        para = hdr[i].paragraphs[0]
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = para.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.name = BODY_FONT
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            mono = val.startswith("`") and val.endswith("`")
            r = para.add_run(val.strip("`"))
            r.font.size = Pt(font_size - (0.5 if mono else 0))
            r.font.name = MONO_FONT if mono else BODY_FONT
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    cap(doc, caption, kind="Table")
    doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
    return t


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
