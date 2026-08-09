"""Builds the AsteroidsFX technical report document. Run this file."""
from build_report import *   # helpers, styles, and the FIG/OUT paths

doc = Document()
style_doc(doc)

# ----------------------------------------------------------------- title page
t = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
t.paragraph_format.space_before = Pt(70)
r = t.add_run("AsteroidsFX")
r.font.size = Pt(30); r.bold = True; r.font.name = BODY_FONT; r.font.color.rgb = ACCENT

s = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
r = s.add_run("A Component-Oriented Game on the Java Platform Module System")
r.font.size = Pt(15); r.font.name = BODY_FONT; r.font.color.rgb = GREY

s = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
r = s.add_run("Designing, implementing and validating runtime-replaceable components")
r.font.size = Pt(12); r.italic = True; r.font.name = BODY_FONT; r.font.color.rgb = GREY

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in [
    ("Student name", "[INSERT YOUR FULL NAME]"),
    ("Student exam number", "[INSERT YOUR EXAM NUMBER]"),
    ("GitHub user", "Fard121"),
    ("GitHub repository", "https://github.com/Fard121/asteroid_game"),
    ("Forked from", "https://github.com/sweat-tek/AsteroidsFX"),
    ("Video demonstration", "[INSERT YOUTUBE LINK]"),
    ("Course", "Component Based Systems, SDU MMMI"),
    ("Contact", "fjama23@student.sdu.dk"),
    ("Date", "9 August 2026"),
]:
    cells = meta.add_row().cells
    cells[0].width = Cm(5.2); cells[1].width = Cm(9.8)
    p0 = cells[0].paragraphs[0]; p0.paragraph_format.space_after = Pt(3)
    r0 = p0.add_run(k); r0.bold = True; r0.font.size = Pt(11.5); r0.font.name = BODY_FONT
    p1 = cells[1].paragraphs[0]; p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(v); r1.font.size = Pt(11.5); r1.font.name = BODY_FONT
    if v.startswith("[INSERT"):
        r1.bold = True; r1.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

P(doc, space_after=0)
n = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
rn = n.add_run("Fields shown in red must be completed before submission.")
rn.italic = True; rn.font.size = Pt(10); rn.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
pagebreak(doc)

# ------------------------------------------------------------------------ ToC
H(doc, "Table of Contents", 1)
import json as _json, os as _os
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

_toc_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "toc.json")
if _os.path.exists(_toc_path):
    for _title, _lvl, _pg in _json.load(open(_toc_path)):
        _p = doc.add_paragraph()
        _p.paragraph_format.space_after = Pt(4)
        _p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _p.paragraph_format.left_indent = Cm(0.0 if _lvl == 1 else 0.8)
        _p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        _r = _p.add_run(_title + "\t" + str(_pg))
        _r.font.name = BODY_FONT
        _r.font.size = Pt(11.5 if _lvl == 1 else 11)
        _r.bold = (_lvl == 1)
else:
    ptoc = doc.add_paragraph()
    add_field(ptoc, 'TOC ' + chr(92) + 'o "1-2" ' + chr(92) + 'h')
pagebreak(doc)

# -------------------------------------------------------------------- abstract
H(doc, "Abstract", 1)
P(doc,
  "This report documents the design, implementation and validation of AsteroidsFX, a "
  "component-oriented reimplementation of the arcade game Asteroids built on the Java Platform "
  "Module System (JPMS), JavaFX and the Spring container. The engineering problem addressed is not "
  "the game itself but the coupling between the parts that make it. In a conventional implementation "
  "the main loop must name every gameplay subsystem it drives, so no subsystem can be added, removed "
  "or replaced without recompiling and relinking that loop. The solution presented here inverts that "
  "relationship. Five gameplay components, Player, Enemy, Bullet, Asteroids and Collision, are "
  "expressed purely as providers of service interfaces declared in shared API modules, are discovered "
  "at runtime by ServiceLoader from a plugins directory, and are resolved into a dedicated child "
  "ModuleLayer. The composition root, Core, declares no requires clause for any of them, and the "
  "Spring container injects whatever the module layer yields. Score keeping is externalised further "
  "still, into a separate Spring Boot process reached over HTTP. The result is validated on three "
  "levels: 21 JUnit 5 and Mockito unit tests covering the component contracts, an integration run of "
  "the complete two-process deployment, and an experimental demonstration that a component can be "
  "uninstalled and reinstalled both at runtime and at deployment time without recompiling any source "
  "file.")
P(doc)
rich(doc, [("Keywords: ", "b"),
           ("component-based software engineering, Java Platform Module System, ServiceLoader, "
            "ModuleLayer, whiteboard component model, dependency injection, microservices, operation "
            "contracts, JavaFX.", "i")])
pagebreak(doc)

# ---------------------------------------------------------------- introduction
H(doc, "1  Introduction", 1)

H(doc, "1.1  Background and motivation", 2)
P(doc,
  "Component-based software engineering treats a system as an assembly of independently deployable "
  "units that interact only through explicitly declared interfaces [5]. The promise is substitution: "
  "if a unit's obligations are fully captured by its contract, any implementation honouring that "
  "contract can replace it without disturbing the rest of the system. The difficulty is that ordinary "
  "language mechanisms work against this, because a direct import creates a compile-time dependency, "
  "and a compile-time dependency is precisely what prevents substitution at deployment time.")
P(doc,
  "Java addressed this in two stages. ServiceLoader decouples a consumer from a provider by resolving "
  "implementations at runtime rather than at compile time [3], [4]. The Java Platform Module System, "
  "delivered in Java 9 under JSR 376, added the missing half: reliable configuration and strong "
  "encapsulation, so that the dependency graph is validated when the application starts and a module's "
  "internals are genuinely unreachable unless deliberately exported [1], [2]. Together they allow a "
  "system to be assembled from parts that it never names.")

H(doc, "1.2  Aim and scope", 2)
P(doc,
  "The aim of this project was to build a working Asteroids game whose gameplay subsystems are true "
  "components in the sense above, and then to demonstrate empirically that one of them can be removed "
  "and restored without recompiling the source code. The scope covers the nine laboratory exercises of "
  "the course [12]: the initial JavaFX game, the component and data-oriented refactoring, ServiceLoader "
  "assembly, the three JPMS exercises on module descriptors, services and module layers, Spring "
  "dependency injection in the composition root, the scoring microservice, and unit testing of the "
  "components.")
P(doc,
  "Persistence of high scores, networked multiplayer, and hot reloading of components that were not "
  "present when the process started are outside scope. Section 7 explains why the last of these is a "
  "deliberate boundary rather than an omission.")

H(doc, "1.3  Structure of this report", 2)
P(doc,
  "Section 2 states the functional and non-functional requirements and the interface contracts they "
  "imply. Section 3 analyses the problem as use cases, a domain model and a behavioural model, without "
  "committing to a solution. Section 4 presents the architecture and the operation contracts of every "
  "provided and required interface. Section 5 explains how components are registered, discovered and "
  "accessed in the source code, and where each component model is applied. Section 6 reports "
  "experimental validation, including the dynamic update experiment. Sections 7 and 8 discuss how well "
  "the design met its requirements and conclude.")
pagebreak(doc)

# ---------------------------------------------------------------- requirements
H(doc, "2  Requirements", 1)
P(doc,
  "Requirements are separated into functional requirements, which describe observable behaviour, and "
  "non-functional requirements, which constrain how that behaviour is achieved. For this project the "
  "non-functional requirements carry most of the engineering weight: a monolithic implementation could "
  "satisfy every functional requirement below and still fail the exercise entirely.")

H(doc, "2.1  Functional requirements", 2)
table(doc,
      ["ID", "Requirement", "Component responsible"],
      [
          ["FR1", "The player controls a ship that rotates, thrusts with inertia and wraps around the screen edges", "Player"],
          ["FR2", "The player fires bullets, subject to a cooldown and a bounded number in flight", "Player, Bullet"],
          ["FR3", "Asteroids drift with random headings and wrap around the screen edges", "Asteroids"],
          ["FR4", "An enemy saucer moves randomly and fires at intervals", "Enemy"],
          ["FR5", "Collision between entities is detected using Pythagorean distance", "Collision"],
          ["FR6", "An asteroid hit by a bullet splits into two smaller ones; the smallest is destroyed outright", "Collision, Asteroids"],
          ["FR7", "Ships that collide with an asteroid are destroyed", "Collision"],
          ["FR8", "The player ship and the enemy saucer are destroyed after a number of bullet hits", "Collision"],
          ["FR9", "The score increases by an amount determined by the target destroyed", "Collision"],
          ["FR10", "Waves advance when the field is cleared, raising count and speed; wave 3 is victory", "Asteroids"],
          ["FR11", "The game presents start and pause menus, a help overlay, and game over and victory states", "Core"],
          ["FR12", "The current score is published to an external scoring service", "Core, Scoring"],
      ],
      "Functional requirements and the component responsible for each.",
      widths=[1.5, 9.3, 4.2])

H(doc, "2.2  Non-functional requirements", 2)
table(doc,
      ["ID", "Requirement", "Rationale"],
      [
          ["NFR1", "A gameplay component must be installable and uninstallable without recompiling any source file", "The central objective of the course; verified experimentally in Section 6.4"],
          ["NFR2", "The composition root must not name any gameplay component at compile time", "A requires clause would reintroduce exactly the coupling NFR1 removes"],
          ["NFR3", "Component internals must be unreachable from other components", "Strong encapsulation: only the declared interfaces form the contract [1]"],
          ["NFR4", "The dependency configuration must be validated at start-up, not when a class is first touched", "Reliable configuration: a missing module should fail loudly and early [2]"],
          ["NFR5", "Every component must be unit-testable in isolation, without a running JavaFX stage", "A component that can only be exercised by running the whole system is not a component"],
          ["NFR6", "Failure of the external scoring service must not stop the game loop", "Availability across a process boundary; see Section 4.6"],
      ],
      "Non-functional requirements.",
      widths=[1.5, 6.6, 6.9])

H(doc, "2.3  Mandated components and their interfaces", 2)
P(doc,
  "The exercise requires five components in particular: Player, Enemy, Asteroids, Weapon, realised "
  "here as the Bullet component, and the rendering and drawing component, realised as Core. Player, "
  "Enemy and Weapon must implement provided service interfaces so that they can be updated or removed "
  "without recompilation. Table 3 states the resulting contract obligations.")
table(doc,
      ["Component", "Provides", "Requires (uses)", "Deployed in"],
      [
          ["Core", "Composition root, rendering, HUD and input; provides no service", "The three processing interfaces, indirectly via Common", "`mods-mvn/`"],
          ["Player", "`IGamePluginService`, `IEntityProcessingService`", "`BulletSPI`", "`plugins/`"],
          ["Enemy", "`IGamePluginService`, `IEntityProcessingService`", "`BulletSPI`", "`plugins/`"],
          ["Bullet (Weapon)", "`IGamePluginService`, `IEntityProcessingService`, `BulletSPI`", "none", "`plugins/`"],
          ["Asteroids", "`IGamePluginService`, `IEntityProcessingService`, `IAsteroidSplitter`", "none", "`plugins/`"],
          ["Collision", "`IPostEntityProcessingService`", "`IAsteroidSplitter`", "`plugins/`"],
          ["Scoring", "REST API on port 8081", "none", "separate process"],
      ],
      "Mandated components with their provided and required interfaces.",
      widths=[2.9, 5.3, 4.5, 2.6])
pagebreak(doc)

# -------------------------------------------------------------------- analysis
H(doc, "3  Analysis", 1)
P(doc,
  "This section describes what the system must do and what it must contain. It deliberately avoids "
  "stating how any of it is achieved, which is the subject of Sections 4 and 5.")

H(doc, "3.1  Use cases", 2)
P(doc,
  "Two actors interact with the system. The player drives gameplay. A second actor, the integrator, "
  "represents whoever assembles or reconfigures the deployment. This actor is unusual in a game, and "
  "its presence is a direct consequence of NFR1: use case UC9 exists only because components are meant "
  "to be replaceable, and it is the use case the video demonstration exercises.")
figure(doc, "fig08-usecase.png",
       "Use case model. The integrator actor and UC9 exist because of the component-replaceability "
       "requirement rather than because of gameplay.", 15.0)
table(doc,
      ["Field", "UC9  Install or uninstall a component at runtime"],
      [
          ["Primary actor", "Integrator"],
          ["Precondition", "The application is running and the target component was discovered at start-up"],
          ["Trigger", "The integrator presses key 1, 2 or 3 for the Player, Enemy or Weapon component"],
          ["Main flow",
           "1. Core resolves the key to a named component.  2. If the component is installed, its "
           "IGamePluginService.stop is invoked and its services are removed from the active lists.  "
           "3. If it is uninstalled, start is invoked and its services are added back."],
          ["Postcondition",
           "The component's entities and per-frame behaviour are absent from, or present in, the "
           "running game, and no other component is affected"],
          ["Guarantee", "No source file is recompiled and the JVM is not restarted"],
      ],
      "Expanded description of the use case central to the exercise.",
      widths=[3.2, 12.0])

H(doc, "3.2  Domain model", 2)
P(doc,
  "The game state is data-oriented [14]: entities are plain data carriers and behaviour lives in "
  "separate processing services, rather than each entity owning its own update method. That separation "
  "is what makes behaviour relocatable into a component at all. Entity holds position, rotation, a "
  "collision radius, a polygon outline, a category and hit points. World is a keyed collection of "
  "entities, and GameData aggregates the session state that every component may read.")
figure(doc, "fig03-class-diagram.png",
       "Domain model. Bullet and Asteroid are declared in shared API modules because several components "
       "must recognise them; Player and Enemy are declared inside their own components and are named by "
       "no one else.", 15.5)
P(doc,
  "The distinction visible in the diagram matters. Bullet and Asteroid live in CommonBullet and "
  "CommonAsteroids because more than one component must recognise them: Player and Enemy count bullets "
  "in flight, and Collision asks an asteroid for its size. Player and Enemy, by contrast, are declared "
  "inside their own components, and Collision distinguishes a player from an asteroid using the "
  "EntityCategory enumeration rather than by type. That is why the Collision component compiles with no "
  "knowledge whatsoever of the Player component.")

H(doc, "3.3  Behavioural model", 2)
P(doc,
  "Gameplay is a state machine owned by GameStateManager. Only in the playing and terminal states are "
  "the entity processors executed. In the menu and paused states the world is rendered but not "
  "advanced, which is what allows a menu to be drawn over a frozen field.")
figure(doc, "fig07-state-machine.png",
       "Game state machine as implemented by GameStateManager and interpreted each frame by Core.", 15.5)

H(doc, "3.4  Identified and missing components", 2)
P(doc,
  "The analysis identifies six in-process components, Player, Enemy, Bullet, Asteroids, Collision and "
  "Core, and one out-of-process component, Scoring. It also identifies capabilities that exist in the "
  "behaviour but are not owned by a component of their own:")
bullet(doc, "Spawning and difficulty. Wave progression lives inside AsteroidProcessor and enemy spawn "
            "timing inside EnemyControlSystem. A dedicated spawn service would allow difficulty to be "
            "tuned without modifying either component.")
bullet(doc, "Score persistence. The scoring service holds the score in memory only, so no component is "
            "responsible for retaining a high score between runs.")
bullet(doc, "Audio. SoundManager is a static utility in the shared module rather than a replaceable "
            "provider, so an alternative sound implementation cannot be substituted the way a gameplay "
            "component can.")
P(doc,
  "These are recorded rather than resolved. Each would be a straightforward addition using the service "
  "mechanism already in place, which is itself evidence that the architecture generalises beyond the "
  "components it currently carries.")
pagebreak(doc)

# ---------------------------------------------------------------------- design
H(doc, "4  Design", 1)

H(doc, "4.1  Architectural overview", 2)
P(doc,
  "The system is deployed as two operating system processes. The game process contains two module "
  "layers: the boot layer, holding the composition root and the shared API modules, and a child layer "
  "holding the five gameplay components. The scoring process is a self-contained Spring Boot "
  "application reached over HTTP. The two-layer split inside a single process is the architectural "
  "decision on which every other one depends.")
figure(doc, "fig01-architecture.png",
       "System architecture. Core sits in the boot layer and knows only the interfaces in Common. The "
       "five gameplay components are discovered in plugins/ and resolved into a child ModuleLayer at "
       "start-up. The score crosses a process boundary over HTTP.", 13.5)
P(doc, "Four architectural principles govern the structure:", space_after=4)
bullet(doc, "Composition root. Core owns the game loop but declares no requires clause for any gameplay "
            "component, satisfying NFR2.")
bullet(doc, "Interface segregation. Contracts live in Common, CommonBullet and CommonAsteroids, so a "
            "component depends on a contract rather than on a peer.")
bullet(doc, "Layer isolation. Component jars are resolved into their own configuration and class loader, "
            "so their resolution cannot corrupt the boot layer's module graph.")
bullet(doc, "Externalised state. The authoritative score lives in a separate process, which forces the "
            "interaction to be an explicit and failure-tolerant protocol rather than a method call.")

H(doc, "4.2  Components and connections", 2)
P(doc,
  "Figure 5 shows the assembly in UML component terms. Every connection is a service contract: a "
  "provided interface on one side, a required interface on the other. No connection in the diagram "
  "represents a compile-time dependency between two gameplay components.")
figure(doc, "fig02-component-diagram.png",
       "UML component diagram. Player and Enemy require BulletSPI, which Bullet provides; Collision "
       "requires IAsteroidSplitter, which Asteroids provides. Neither pair shares a compile-time "
       "dependency.", 15.5)
P(doc,
  "Figure 6 shows the same system from the module system's point of view. Solid edges are requires "
  "clauses resolved when the application starts; dashed edges are uses relationships resolved at "
  "runtime by ServiceLoader. The absence of any solid edge from Core to a gameplay module is the "
  "structural proof of NFR2.")
figure(doc, "fig04-module-graph.png",
       "Module dependency graph. Solid edges are requires clauses, dashed edges are runtime service "
       "lookups. Core has no solid edge to any gameplay module.", 15.5)

H(doc, "4.3  Operation contracts", 2)
P(doc,
  "Each interface operation is specified as a contract in the style of the Unified Process [13]: a "
  "precondition the caller must establish and a postcondition the implementation guarantees. These "
  "contracts are written as JavaDoc on the interfaces themselves, so they travel with the code rather "
  "than only with this document.")
table(doc,
      ["Operation", "Precondition", "Postcondition"],
      [
          ["`IGamePluginService.start`",
           "gameData and world are non-null and represent a fresh session; the world may already hold entities added by plugins started earlier, but none of this plugin's own",
           "Entities this plugin owns up front have been added to the world, and plugin-local state needed by its processor has been initialised"],
          ["`IGamePluginService.stop`",
           "start was previously called on this instance with the same gameData and world",
           "Every entity exclusively owned by this plugin has been removed and its resources released; calling stop without a prior start is a no-op, not an error"],
          ["`IEntityProcessingService.process`",
           "Called once per frame in no guaranteed order relative to other implementations, so it may not assume any peer has already run this frame",
           "Every entity this component owns has advanced exactly one frame and any entities it spawns have been added; it never removes another component's entities"],
          ["`IPostEntityProcessingService.process`",
           "Every IEntityProcessingService for this frame has completed, so all entity positions are final",
           "Cross-entity effects are fully applied for the frame; the implementation must snapshot the world before mutating it so that no pair is skipped or processed twice"],
          ["`BulletSPI.createBullet`",
           "The shooter is non-null and has a valid position and rotation",
           "A fully initialised bullet is returned but not added to the world, or empty if the weapon component is not currently installed; adding it is the caller's responsibility"],
          ["`IAsteroidSplitter.createSplitAsteroid`",
           "The entity is an asteroid that the caller has removed or is about to remove; this operation does not remove it",
           "Two smaller asteroids have been added at the destroyed asteroid's position, or none if it was already the smallest size"],
      ],
      "Operation contracts for every provided and required interface.",
      widths=[3.5, 5.9, 5.9], font_size=9)

H(doc, "4.4  The two-phase frame", 2)
P(doc,
  "The precondition on IEntityProcessingService.process, that implementations run in no guaranteed "
  "order, is the reason collision detection cannot be one of them. If collision ran as an ordinary "
  "processor it could observe some entities at their new positions and others at their old ones, and "
  "the outcome would depend on the order in which ServiceLoader happened to return providers. Dividing "
  "the frame into an entity phase and a post-entity phase makes the ordering guarantee explicit in the "
  "type system: any logic that requires settled positions must implement the second interface, and the "
  "compiler, rather than a comment, enforces the distinction.")
figure(doc, "fig06-gameloop-flow.png",
       "Per-frame control flow. The simulation phases are skipped in the menu and paused states, and "
       "score synchronisation deliberately runs last.", 15.5)

H(doc, "4.5  Module layers and split packages", 2)
P(doc,
  "Two modules that export the same package cannot be defined under one class loader; the module system "
  "rejects the attempt with a LayerInstantiationException. This was reproduced deliberately in a "
  "minimal two-module demonstration kept under docs/jpms-lab3-demo, and the captured behaviour is "
  "instructive. Configuration.resolve alone does not detect the clash, because it resolves only the "
  "requires graph. The failure appears when the modules are actually defined together, which is exactly "
  "what happens when everything sits on one flat module path.")
P(doc,
  "The resolution is to give each conflicting module its own layer, and therefore its own loader. "
  "AsteroidsFX applies the same mechanism for a related reason: because the plugin set is resolved as "
  "its own isolated configuration, a package clash between two components fails loudly and specifically "
  "when ServiceLocator is constructed, instead of silently corrupting the boot layer's module graph.")

H(doc, "4.6  The scoring boundary", 2)
P(doc,
  "Moving the score out of process turns a field access into a network call, and a network call can "
  "fail. NFR6 requires that this failure be contained, and two design decisions achieve it. First, "
  "synchronisation runs as the last statement of the frame, after rendering and input handling have "
  "already completed, so an exception cannot abort the remainder of the frame. Second, the client "
  "publishes only on an actual change of score, and after a failure backs off for roughly three seconds "
  "rather than retrying every frame.")
figure(doc, "fig12-seq-score-sync.png",
       "Score synchronisation including the failure path. The game continues at full frame rate whether "
       "or not the scoring service is reachable.", 13.5)
P(doc,
  "This design was not arrived at analytically. An earlier version placed the call before rendering, "
  "and when the scoring service was stopped the uncaught exception aborted the rest of the frame: the "
  "screen stopped redrawing and input stopped responding, which presented exactly as a frozen game. The "
  "ordering is therefore a corrective response to an observed defect, and is documented as such in the "
  "repository.")
pagebreak(doc)

# ------------------------------------------- sections 5 to 9 and appendices
import report_content2
report_content2.build(doc, H, P, rich, bullet, code, figure, table, pagebreak, cap,
                      WD_ALIGN_PARAGRAPH, Pt, GREY)

# ---------------------------------------------------------------- page numbers
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _El

for section in doc.sections:
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer_p, "PAGE")
    for run in footer_p.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = GREY

doc.save(OUT)
print("SAVED", OUT)
print("figures:", fig_no["n"], "tables:", tab_no["n"])
